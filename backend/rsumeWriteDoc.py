from fastapi.responses import StreamingResponse
from io import BytesIO
from docx import Document

@app.post("/generate_resume_docx/")
async def generate_resume_docx(data: ResumeRequest):
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return {"error": "OPENAI_API_KEY environment variable not set."}

    prompt = (
        "You are a professional resume writer. "
        "Given the following job description and candidate profile, generate a customized resume. "
        "IMPORTANT INSTRUCTIONS:\n"
        "1. Start the resume with the candidate's NAME and EMAIL.\n"
        "2. Include a section for JOB TITLE or the job being applied for.\n"
        "3. Format ALL job responsibilities and achievements as bullet points (start each with '- ').\n"
        "4. HIGHLIGHT ALL TECHNICAL TERMS in Bold (such as programming languages, frameworks, tools, and certifications) BY MAKING THEM UPPERCASE.\n"
        "5. Do NOT use placeholder text like [Your Name] or [Your Email Address]; use the provided name and email.\n"
        "Example:\n"
        "Name: JOHN DOE\n"
        "Email: JOHN.DOE@EMAIL.COM\n"
        "Job Title: BACKEND DEVELOPER\n"
        "Experience:\n"
        "- Developed scalable APIs using PYTHON and FASTAPI.\n"
        "- Managed cloud infrastructure on AWS.\n"
        "- Improved database performance using SQL optimization techniques.\n\n"
        f"Name: {data.name}\n"
        f"Email: {data.email}\n"
        f"Job Title: {data.job_description}\n"
        f"Candidate Profile:\n{data.experience}\n\n"
        "Customized Resume:"
    )

    client = openai.OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1000,
        temperature=0.7,
    )

    resume_text = response.choices[0].message.content.strip()

    # Create Word document
    doc = Document()
    for line in resume_text.split('\n'):
        if line.startswith('- '):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)

    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=resume_{data.name.replace(' ', '_')}.docx"}
    )